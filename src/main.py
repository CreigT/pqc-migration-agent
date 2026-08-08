#!/usr/bin/env python3
"""Post-Quantum Cryptography migration agent.

This is a local, deterministic agent for scanning PDF, DOCX, and TXT documents.
It extracts real document text, finds cryptographic references with explicit
rules, and writes JSON plus Markdown reports from actual evidence only.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import re
import sys
from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Iterable


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}
RISK_ORDER = {"critical": 0, "high": 1, "medium": 2, "low": 3}
EXIT_OK = 0
EXIT_RUNTIME_ERROR = 1
EXIT_USAGE_ERROR = 2
EXIT_FINDINGS_PRESENT = 3


@dataclass(frozen=True)
class DocumentBlock:
    text: str
    source_type: str
    page: int | None = None
    line: int | None = None
    paragraph: int | None = None
    table: int | None = None
    row: int | None = None
    cell: int | None = None


@dataclass(frozen=True)
class CryptoRule:
    rule_id: str
    name: str
    category: str
    risk: str
    post_quantum_relevant: bool
    description: str
    recommendation: str
    patterns: tuple[re.Pattern[str], ...] = field(default_factory=tuple)


@dataclass(frozen=True)
class AgentOptions:
    context_chars: int
    recursive: bool
    include_clean_documents: bool


def compile_patterns(*patterns: str) -> tuple[re.Pattern[str], ...]:
    return tuple(re.compile(pattern, re.IGNORECASE) for pattern in patterns)


CRYPTO_RULES: tuple[CryptoRule, ...] = (
    CryptoRule(
        rule_id="rsa-weak-key",
        name="Weak RSA key size",
        category="public-key encryption/signature",
        risk="critical",
        post_quantum_relevant=True,
        description="RSA keys below 2048 bits are already considered weak, and RSA is not post-quantum safe.",
        recommendation="Replace RSA-1024 and RSA-1536 immediately. For post-quantum migration, move encryption and key-establishment use cases to ML-KEM or an approved hybrid profile, and move signature use cases to ML-DSA or SLH-DSA.",
        patterns=compile_patterns(
            r"\bRSA[\s_-]*(?:1024|1,024|1536|1,536)\b",
            r"\b(?:1024|1,024|1536|1,536)[\s_-]*(?:bit|bits)?[\s_-]*RSA\b",
            r"\bRSA\b.{0,40}\b(?:key\s*(?:size|length)|modulus)\b.{0,24}\b(?:1024|1,024|1536|1,536)\b",
        ),
    ),
    CryptoRule(
        rule_id="rsa-general",
        name="RSA",
        category="public-key encryption/signature",
        risk="high",
        post_quantum_relevant=True,
        description="RSA public-key cryptography is vulnerable to future cryptographically relevant quantum computers.",
        recommendation="Inventory each RSA dependency by use case. Plan migration to ML-KEM for key establishment or encryption, and ML-DSA or SLH-DSA for signatures. Use hybrid modes where ecosystem constraints require a transition period.",
        patterns=compile_patterns(r"\bRSA\b(?![\s_-]*(?:1024|1,024|1536|1,536))"),
    ),
    CryptoRule(
        rule_id="dsa",
        name="DSA",
        category="digital signature",
        risk="high",
        post_quantum_relevant=True,
        description="DSA is vulnerable to quantum attacks and is not appropriate for new designs.",
        recommendation="Replace DSA signatures with ML-DSA or SLH-DSA after validating library, protocol, certificate, and HSM support.",
        patterns=compile_patterns(r"\bDSA\b", r"\bDigital\s+Signature\s+Algorithm\b"),
    ),
    CryptoRule(
        rule_id="ecc",
        name="ECC / elliptic curve cryptography",
        category="public-key encryption/signature/key agreement",
        risk="high",
        post_quantum_relevant=True,
        description="Elliptic-curve public-key cryptography is vulnerable to future quantum attacks.",
        recommendation="Map ECC usage by purpose. Replace ECDH key agreement with ML-KEM or a vetted hybrid KEM, and replace ECDSA signatures with ML-DSA or SLH-DSA.",
        patterns=compile_patterns(
            r"\bECC\b",
            r"\belliptic[-\s]+curve\b",
            r"\bsecp(?:224|256|384|521)r1\b",
            r"\bprime256v1\b",
            r"\bP-(?:224|256|384|521)\b",
            r"\bnistp(?:224|256|384|521)\b",
        ),
    ),
    CryptoRule(
        rule_id="ecdsa",
        name="ECDSA",
        category="digital signature",
        risk="high",
        post_quantum_relevant=True,
        description="ECDSA signatures are vulnerable to future quantum attacks.",
        recommendation="Plan migration from ECDSA to ML-DSA or SLH-DSA. Where compatibility requires transition, evaluate dual-signature or hybrid certificate approaches.",
        patterns=compile_patterns(r"\bECDSA\b"),
    ),
    CryptoRule(
        rule_id="ecdh",
        name="ECDH / ECDHE",
        category="key agreement",
        risk="high",
        post_quantum_relevant=True,
        description="ECDH and ECDHE key agreement are vulnerable to future quantum attacks.",
        recommendation="Replace ECDH/ECDHE with ML-KEM or an approved hybrid key-establishment profile during migration.",
        patterns=compile_patterns(r"\bECDH\b", r"\bECDHE\b"),
    ),
    CryptoRule(
        rule_id="dh-weak",
        name="Weak finite-field Diffie-Hellman",
        category="key agreement",
        risk="high",
        post_quantum_relevant=True,
        description="Small finite-field Diffie-Hellman groups are weak today, and Diffie-Hellman is not post-quantum safe.",
        recommendation="Retire 1024-bit finite-field Diffie-Hellman groups. Replace long-term key agreement with ML-KEM or a vetted hybrid post-quantum profile.",
        patterns=compile_patterns(
            r"\b(?:DH|DHE|Diffie[-\s]*Hellman)\b.{0,40}\b(?:1024|1,024)\b",
            r"\b(?:1024|1,024)[\s_-]*(?:bit|bits)?[\s_-]*(?:DH|DHE|Diffie[-\s]*Hellman)\b",
            r"\b(?:Oakley\s+Group\s+2|modp1024|Group\s+2)\b",
        ),
    ),
    CryptoRule(
        rule_id="sha1",
        name="SHA-1",
        category="hash/signature digest",
        risk="high",
        post_quantum_relevant=False,
        description="SHA-1 has practical collision attacks and should not be used for signatures, certificates, integrity, or new systems.",
        recommendation="Replace SHA-1 with SHA-256, SHA-384, SHA-512, or SHA-3 according to the protocol and compliance requirement.",
        patterns=compile_patterns(r"\bSHA[\s_-]?1\b", r"\bSHA1\b"),
    ),
    CryptoRule(
        rule_id="md5",
        name="MD5",
        category="hash",
        risk="critical",
        post_quantum_relevant=False,
        description="MD5 is cryptographically broken and unsuitable for integrity, signatures, certificates, or password storage.",
        recommendation="Replace MD5 with SHA-256, SHA-384, SHA-512, or SHA-3 for general hashing. For password storage, use Argon2id, bcrypt, scrypt, or PBKDF2 with current parameters.",
        patterns=compile_patterns(r"\bMD5\b"),
    ),
    CryptoRule(
        rule_id="des",
        name="DES",
        category="symmetric encryption",
        risk="critical",
        post_quantum_relevant=False,
        description="DES has an obsolete 56-bit key size and is brute-forceable.",
        recommendation="Replace DES with AES-GCM, ChaCha20-Poly1305, or another approved authenticated encryption mode.",
        patterns=compile_patterns(r"(?<!3)\bDES\b", r"\bData\s+Encryption\s+Standard\b"),
    ),
    CryptoRule(
        rule_id="3des",
        name="3DES / Triple DES",
        category="symmetric encryption",
        risk="high",
        post_quantum_relevant=False,
        description="3DES is deprecated for most uses and has a small block size that creates practical risk at scale.",
        recommendation="Migrate 3DES or Triple DES to AES-GCM, AES-CTR with authentication, or ChaCha20-Poly1305. Use AES-CBC only where legacy constraints require it and pair it with strong integrity protection.",
        patterns=compile_patterns(r"\b3DES\b", r"\bTriple[\s_-]*DES\b", r"\bTDEA\b"),
    ),
    CryptoRule(
        rule_id="rc4",
        name="RC4",
        category="stream cipher",
        risk="critical",
        post_quantum_relevant=False,
        description="RC4 has severe statistical biases and is prohibited in modern secure protocols.",
        recommendation="Remove RC4 support and use AES-GCM or ChaCha20-Poly1305.",
        patterns=compile_patterns(r"\bRC4\b", r"\bARCFOUR\b", r"\bARC4\b"),
    ),
)


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as file:
        for chunk in iter(lambda: file.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def file_metadata(path: Path) -> dict[str, Any]:
    stat = path.stat()
    return {
        "path": str(path.resolve()),
        "file_name": path.name,
        "file_type": path.suffix.lower().lstrip("."),
        "size_bytes": stat.st_size,
        "sha256": sha256_file(path),
        "modified_at": datetime.fromtimestamp(stat.st_mtime, timezone.utc).isoformat(),
    }


def read_txt(path: Path) -> tuple[list[DocumentBlock], dict[str, Any]]:
    text = path.read_text(encoding="utf-8-sig", errors="replace")
    lines = text.splitlines()
    blocks = [DocumentBlock(text=line, source_type="txt", line=index) for index, line in enumerate(lines, start=1)]
    return blocks, {"line_count": len(lines), "character_count": len(text)}


def read_pdf(path: Path) -> tuple[list[DocumentBlock], dict[str, Any]]:
    try:
        import pdfplumber
    except ImportError as exc:
        raise RuntimeError("PDF input requires pdfplumber. Install it with: pip install pdfplumber") from exc

    blocks: list[DocumentBlock] = []
    extracted_characters = 0
    with pdfplumber.open(path) as pdf:
        page_count = len(pdf.pages)
        for page_index, page in enumerate(pdf.pages, start=1):
            page_text = page.extract_text() or ""
            extracted_characters += len(page_text)
            for line_index, line in enumerate(page_text.splitlines(), start=1):
                blocks.append(DocumentBlock(text=line, source_type="pdf", page=page_index, line=line_index))
    return blocks, {"page_count": page_count, "extracted_character_count": extracted_characters}


def read_docx(path: Path) -> tuple[list[DocumentBlock], dict[str, Any]]:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("DOCX input requires python-docx. Install it with: pip install python-docx") from exc

    document = Document(path)
    blocks: list[DocumentBlock] = []
    paragraph_count = len(document.paragraphs)
    table_count = len(document.tables)

    for index, paragraph in enumerate(document.paragraphs, start=1):
        text = paragraph.text.strip()
        if text:
            blocks.append(DocumentBlock(text=text, source_type="docx", paragraph=index))

    table_cell_count = 0
    for table_index, table in enumerate(document.tables, start=1):
        for row_index, row in enumerate(table.rows, start=1):
            for cell_index, cell in enumerate(row.cells, start=1):
                table_cell_count += 1
                text = " ".join(paragraph.text.strip() for paragraph in cell.paragraphs if paragraph.text.strip())
                if text:
                    blocks.append(
                        DocumentBlock(
                            text=text,
                            source_type="docx_table_cell",
                            table=table_index,
                            row=row_index,
                            cell=cell_index,
                        )
                    )

    return blocks, {"paragraph_count": paragraph_count, "table_count": table_count, "table_cell_count": table_cell_count}


def read_document(path: Path) -> tuple[list[DocumentBlock], dict[str, Any]]:
    extension = path.suffix.lower()
    if extension == ".txt":
        return read_txt(path)
    if extension == ".pdf":
        return read_pdf(path)
    if extension == ".docx":
        return read_docx(path)
    raise ValueError(f"Unsupported input type '{extension}'. Supported types: {', '.join(sorted(SUPPORTED_EXTENSIONS))}")


def block_location(block: DocumentBlock) -> dict[str, Any]:
    location: dict[str, Any] = {"source_type": block.source_type}
    if block.page is not None:
        location["page"] = block.page
    if block.line is not None:
        location["line"] = block.line
    if block.paragraph is not None:
        location["paragraph"] = block.paragraph
    if block.table is not None:
        location["table"] = block.table
    if block.row is not None:
        location["row"] = block.row
    if block.cell is not None:
        location["cell"] = block.cell
    return location


def location_label(location: dict[str, Any]) -> str:
    parts = [str(location.get("source_type", "document"))]
    for key in ("page", "line", "paragraph", "table", "row", "cell"):
        if key in location:
            parts.append(f"{key} {location[key]}")
    return ", ".join(parts)


def excerpt_for_match(text: str, start: int, end: int, context_chars: int) -> str:
    left = max(0, start - context_chars)
    right = min(len(text), end + context_chars)
    prefix = "..." if left > 0 else ""
    suffix = "..." if right < len(text) else ""
    return f"{prefix}{text[left:right].strip()}{suffix}"


def find_crypto_references(blocks: Iterable[DocumentBlock], context_chars: int) -> list[dict[str, Any]]:
    findings: list[dict[str, Any]] = []
    seen: set[tuple[str, tuple[tuple[str, Any], ...], int, int, str]] = set()

    for block in blocks:
        if not block.text.strip():
            continue

        location = block_location(block)
        location_key = tuple(sorted(location.items()))
        for rule in CRYPTO_RULES:
            for pattern in rule.patterns:
                for match in pattern.finditer(block.text):
                    match_text = match.group(0)
                    dedupe_key = (rule.rule_id, location_key, match.start(), match.end(), match_text.lower())
                    if dedupe_key in seen:
                        continue
                    seen.add(dedupe_key)

                    findings.append(
                        {
                            "rule_id": rule.rule_id,
                            "algorithm": rule.name,
                            "category": rule.category,
                            "risk": rule.risk,
                            "post_quantum_relevant": rule.post_quantum_relevant,
                            "description": rule.description,
                            "recommendation": rule.recommendation,
                            "matched_text": match_text,
                            "match_start": match.start(),
                            "match_end": match.end(),
                            "location": location,
                            "excerpt": excerpt_for_match(block.text, match.start(), match.end(), context_chars),
                        }
                    )

    findings.sort(key=lambda item: (RISK_ORDER.get(item["risk"], 99), item["algorithm"], str(item["location"])))
    return findings


def risk_counts(findings: Iterable[dict[str, Any]]) -> dict[str, int]:
    counts = {"critical": 0, "high": 0, "medium": 0, "low": 0}
    for finding in findings:
        counts[finding["risk"]] = counts.get(finding["risk"], 0) + 1
    return counts


def summarize_rules(findings: list[dict[str, Any]]) -> list[dict[str, Any]]:
    by_rule: dict[str, dict[str, Any]] = {}
    for finding in findings:
        entry = by_rule.setdefault(
            finding["rule_id"],
            {
                "rule_id": finding["rule_id"],
                "algorithm": finding["algorithm"],
                "risk": finding["risk"],
                "category": finding["category"],
                "post_quantum_relevant": finding["post_quantum_relevant"],
                "finding_count": 0,
                "recommendation": finding["recommendation"],
            },
        )
        entry["finding_count"] += 1
    return sorted(by_rule.values(), key=lambda item: (RISK_ORDER.get(item["risk"], 99), item["algorithm"]))


class PQCMigrationAgent:
    """Coordinates discovery, parsing, analysis, and report assembly."""

    def __init__(self, options: AgentOptions) -> None:
        self.options = options

    def discover_targets(self, target: Path) -> tuple[list[Path], list[dict[str, Any]]]:
        unsupported: list[dict[str, Any]] = []
        if target.is_file():
            if target.suffix.lower() in SUPPORTED_EXTENSIONS:
                return [target], unsupported
            unsupported.append({"path": str(target.resolve()), "reason": "unsupported_file_type"})
            return [], unsupported

        pattern = "**/*" if self.options.recursive else "*"
        targets: list[Path] = []
        for path in sorted(target.glob(pattern)):
            if not path.is_file():
                continue
            if path.suffix.lower() in SUPPORTED_EXTENSIONS:
                targets.append(path)
            else:
                unsupported.append({"path": str(path.resolve()), "reason": "unsupported_file_type"})
        return targets, unsupported

    def analyze_file(self, path: Path) -> dict[str, Any]:
        metadata = file_metadata(path)
        try:
            blocks, extraction = read_document(path)
            findings = find_crypto_references(blocks, self.options.context_chars)
            status = "completed"
            error = None
        except Exception as exc:
            blocks = []
            extraction = {}
            findings = []
            status = "failed"
            error = str(exc)

        report: dict[str, Any] = {
            "status": status,
            "input": metadata,
            "extraction": {
                **extraction,
                "parsed_blocks": len(blocks),
                "non_empty_blocks": sum(1 for block in blocks if block.text.strip()),
            },
            "summary": {
                "total_findings": len(findings),
                "risk_counts": risk_counts(findings),
                "post_quantum_relevant_findings": sum(1 for finding in findings if finding["post_quantum_relevant"]),
                "rules_triggered": summarize_rules(findings),
            },
            "findings": findings,
        }
        if error is not None:
            report["error"] = error
        return report

    def run(self, target: Path) -> dict[str, Any]:
        started_at = datetime.now(timezone.utc)
        targets, unsupported = self.discover_targets(target)
        document_reports = [self.analyze_file(path) for path in targets]

        if not self.options.include_clean_documents:
            visible_reports = [
                report for report in document_reports if report["status"] != "completed" or report["summary"]["total_findings"] > 0
            ]
        else:
            visible_reports = document_reports

        completed_reports = [report for report in document_reports if report["status"] == "completed"]
        failed_reports = [report for report in document_reports if report["status"] == "failed"]
        all_findings = [
            finding
            for report in completed_reports
            for finding in report["findings"]
        ]
        finished_at = datetime.now(timezone.utc)

        return {
            "agent": {
                "name": "Post-Quantum Cryptography Migration Agent",
                "version": "2.0.0",
                "started_at": started_at.isoformat(),
                "finished_at": finished_at.isoformat(),
                "duration_seconds": round((finished_at - started_at).total_seconds(), 6),
                "supported_file_types": sorted(SUPPORTED_EXTENSIONS),
                "rules_loaded": len(CRYPTO_RULES),
            },
            "target": {
                "path": str(target.resolve()),
                "mode": "directory" if target.is_dir() else "file",
                "recursive": self.options.recursive if target.is_dir() else False,
            },
            "processing": {
                "documents_discovered": len(targets),
                "documents_completed": len(completed_reports),
                "documents_failed": len(failed_reports),
                "unsupported_files_seen": len(unsupported),
                "clean_documents_omitted_from_report": len(document_reports) - len(visible_reports),
            },
            "summary": {
                "total_findings": len(all_findings),
                "risk_counts": risk_counts(all_findings),
                "post_quantum_relevant_findings": sum(1 for finding in all_findings if finding["post_quantum_relevant"]),
                "rules_triggered": summarize_rules(all_findings),
                "affected_documents": sum(1 for report in completed_reports if report["summary"]["total_findings"] > 0),
            },
            "documents": visible_reports,
            "unsupported_files": unsupported,
        }


def markdown_summary(report: dict[str, Any]) -> str:
    summary = report["summary"]
    processing = report["processing"]
    target = report["target"]

    lines = [
        "# Post-Quantum Cryptography Migration Agent Report",
        "",
        f"- Target: `{target['path']}`",
        f"- Mode: `{target['mode']}`",
        f"- Documents discovered: {processing['documents_discovered']}",
        f"- Documents completed: {processing['documents_completed']}",
        f"- Documents failed: {processing['documents_failed']}",
        f"- Unsupported files seen: {processing['unsupported_files_seen']}",
        f"- Total findings: {summary['total_findings']}",
        f"- Affected documents: {summary['affected_documents']}",
        f"- Post-quantum relevant findings: {summary['post_quantum_relevant_findings']}",
        "",
        "## Risk Counts",
        "",
    ]

    for risk in ("critical", "high", "medium", "low"):
        lines.append(f"- {risk.title()}: {summary['risk_counts'].get(risk, 0)}")

    lines.extend(["", "## Triggered Rules", ""])
    if not summary["rules_triggered"]:
        lines.append("No configured cryptography rules were triggered.")
    else:
        for rule in summary["rules_triggered"]:
            lines.append(
                f"- {rule['algorithm']} ({rule['risk'].title()}): {rule['finding_count']} finding(s), "
                f"category `{rule['category']}`"
            )

    lines.extend(["", "## Document Findings", ""])
    if not report["documents"]:
        lines.append("No reportable document findings.")
    else:
        for document in report["documents"]:
            input_info = document["input"]
            lines.extend(
                [
                    f"### {input_info['file_name']}",
                    "",
                    f"- Path: `{input_info['path']}`",
                    f"- SHA-256: `{input_info['sha256']}`",
                    f"- Size bytes: {input_info['size_bytes']}",
                    f"- Status: {document['status']}",
                    f"- Parsed blocks: {document['extraction'].get('parsed_blocks', 0)}",
                    f"- Findings: {document['summary']['total_findings']}",
                    "",
                ]
            )
            if document["status"] == "failed":
                lines.extend([f"Error: {document.get('error', 'unknown error')}", ""])
                continue
            if not document["findings"]:
                lines.extend(["No vulnerable or migration-relevant cryptography references were detected.", ""])
                continue
            for index, finding in enumerate(document["findings"], start=1):
                lines.extend(
                    [
                        f"{index}. **{finding['algorithm']}** ({finding['risk'].title()})",
                        f"   - Location: {location_label(finding['location'])}",
                        f"   - Matched text: `{finding['matched_text']}`",
                        f"   - Evidence: {finding['excerpt']}",
                        f"   - Recommendation: {finding['recommendation']}",
                    ]
                )
            lines.append("")

    lines.extend(["## Migration Actions", ""])
    if summary["total_findings"] == 0:
        lines.append("No migration actions were generated because the configured rules did not find vulnerable or migration-relevant cryptography in the parsed text.")
    else:
        lines.extend(
            [
                "1. Validate each finding against the source system, cryptographic library, certificate inventory, HSM configuration, and runtime configuration.",
                "2. Prioritize critical and high-risk findings before medium and low-risk findings.",
                "3. Replace quantum-vulnerable public-key encryption or key-establishment dependencies with ML-KEM or vetted hybrid profiles.",
                "4. Replace quantum-vulnerable signature dependencies with ML-DSA or SLH-DSA after protocol and ecosystem validation.",
                "5. Replace broken or deprecated non-PQC primitives such as MD5, SHA-1, DES, 3DES, and RC4 with approved modern algorithms.",
                "6. Re-scan the updated documentation and configuration evidence after remediation.",
            ]
        )

    if report["unsupported_files"]:
        lines.extend(["", "## Unsupported Files", ""])
        for item in report["unsupported_files"]:
            lines.append(f"- `{item['path']}`: {item['reason']}")

    lines.extend(
        [
            "",
            "## Evidence Policy",
            "",
            "This agent reports only findings produced from parsed document text and explicit rules. It does not create confidence scores, synthetic assets, or mock data.",
            "",
        ]
    )
    return "\n".join(lines)


def default_output_paths(target: Path, output_dir: Path | None) -> tuple[Path, Path]:
    target_dir = output_dir if output_dir is not None else (target if target.is_dir() else target.parent)
    base_name = target.resolve().name if target.is_dir() else target.stem
    return target_dir / f"{base_name}.pqc_report.json", target_dir / f"{base_name}.pqc_summary.md"


def parse_args(argv: list[str]) -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Operate a real local PQC migration agent over PDF, DOCX, and TXT documents."
    )
    parser.add_argument("target", help="Path to a PDF, DOCX, TXT file, or a directory containing supported files.")
    parser.add_argument("--json-output", help="Path for the JSON report. Defaults beside the input target.")
    parser.add_argument("--markdown-output", help="Path for the Markdown summary. Defaults beside the input target.")
    parser.add_argument("--output-dir", help="Directory for default report outputs.")
    parser.add_argument("--recursive", action="store_true", help="Recursively scan directories.")
    parser.add_argument(
        "--include-clean-documents",
        action="store_true",
        help="Include successfully scanned documents with zero findings in the JSON and Markdown report.",
    )
    parser.add_argument(
        "--context-chars",
        type=int,
        default=80,
        help="Number of surrounding characters to include in evidence excerpts. Default: 80.",
    )
    parser.add_argument("--print-summary", action="store_true", help="Print the Markdown summary after writing files.")
    parser.add_argument(
        "--fail-on-findings",
        action="store_true",
        help=f"Exit with code {EXIT_FINDINGS_PRESENT} when one or more findings are present.",
    )
    return parser.parse_args(argv)


def run(argv: list[str]) -> int:
    args = parse_args(argv)
    target = Path(args.target).expanduser()
    if not target.exists():
        print(f"Target does not exist: {target}", file=sys.stderr)
        return EXIT_USAGE_ERROR
    if not target.is_file() and not target.is_dir():
        print(f"Target is neither a file nor a directory: {target}", file=sys.stderr)
        return EXIT_USAGE_ERROR

    output_dir = Path(args.output_dir).expanduser() if args.output_dir else None
    if output_dir is not None:
        output_dir.mkdir(parents=True, exist_ok=True)

    default_json_path, default_markdown_path = default_output_paths(target, output_dir)
    json_path = Path(args.json_output).expanduser() if args.json_output else default_json_path
    markdown_path = Path(args.markdown_output).expanduser() if args.markdown_output else default_markdown_path
    json_path.parent.mkdir(parents=True, exist_ok=True)
    markdown_path.parent.mkdir(parents=True, exist_ok=True)

    options = AgentOptions(
        context_chars=max(0, args.context_chars),
        recursive=args.recursive,
        include_clean_documents=args.include_clean_documents,
    )
    agent = PQCMigrationAgent(options)

    try:
        report = agent.run(target)
        markdown = markdown_summary(report)
        json_path.write_text(json.dumps(report, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")
        markdown_path.write_text(markdown, encoding="utf-8")
    except Exception as exc:
        print(f"Agent failed: {exc}", file=sys.stderr)
        return EXIT_RUNTIME_ERROR

    print(f"JSON report written: {json_path.resolve()}")
    print(f"Markdown summary written: {markdown_path.resolve()}")
    print(f"Documents completed: {report['processing']['documents_completed']}")
    print(f"Findings: {report['summary']['total_findings']}")

    if args.print_summary:
        print()
        print(markdown)

    if report["processing"]["documents_failed"] > 0:
        return EXIT_RUNTIME_ERROR
    if args.fail_on_findings and report["summary"]["total_findings"] > 0:
        return EXIT_FINDINGS_PRESENT
    return EXIT_OK


if __name__ == "__main__":
    raise SystemExit(run(sys.argv[1:]))
